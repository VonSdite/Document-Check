from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
import fitz
from openpyxl import load_workbook
from pypdf import PdfReader
import xlrd


ALLOWED_EXTENSIONS = {"docx", "pdf", "txt", "md", "html", "xlsx", "xlsm", "xls"}


class DocumentReadError(Exception):
    pass


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extension_of(filename: str) -> str:
    return filename.rsplit(".", 1)[1].lower()


def extract_text(path: Path, file_type: str) -> str:
    try:
        if file_type == "docx":
            return _extract_docx(path)
        if file_type == "pdf":
            return _extract_pdf(path)
        if file_type in {"txt", "md"}:
            return _read_text(path)
        if file_type == "html":
            return _extract_html(path)
        if file_type in {"xlsx", "xlsm"}:
            return _extract_openpyxl_workbook(path)
        if file_type == "xls":
            return _extract_xls(path)
    except Exception as exc:
        raise DocumentReadError(str(exc)) from exc
    raise DocumentReadError(f"不支持的文件类型：{file_type}")


def format_document_text(filename: str, text: str) -> str:
    text = str(text or "").strip()
    if not text:
        return ""
    name = Path(str(filename or "")).name.strip()
    if not name:
        return text
    return f"file: {name}\n\n{text}"


def _extract_docx(path: Path) -> str:
    document = Document(str(path))
    parts = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    layout_document = fitz.open(str(path))
    try:
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if index <= layout_document.page_count:
                text = _normalize_pdf_overlapping_spaces(text, layout_document[index - 1])
            if text.strip():
                pages.append(f"[第{index}页]\n{text.strip()}")
        return "\n\n".join(pages)
    finally:
        layout_document.close()


def _normalize_pdf_overlapping_spaces(text: str, page) -> str:
    """Remove encoded spaces that have no visible advance on the PDF page."""
    if not text:
        return text
    try:
        raw_page = page.get_text("rawdict")
    except Exception:
        return text

    source_counts = {}
    replacements = {}
    for block in raw_page.get("blocks", []):
        for line in block.get("lines", []):
            source, normalized = _pdf_line_text_variants(line)
            if not source:
                continue
            occurrence = source_counts.get(source, 0) + 1
            source_counts[source] = occurrence
            if normalized != source:
                replacements[(source, occurrence)] = normalized

    if not replacements:
        return text

    text_counts = {}
    normalized_lines = []
    for line in text.splitlines(keepends=True):
        content = line.rstrip("\r\n")
        line_ending = line[len(content) :]
        occurrence = text_counts.get(content, 0) + 1
        text_counts[content] = occurrence
        normalized_lines.append(replacements.get((content, occurrence), content) + line_ending)
    return "".join(normalized_lines)


def _pdf_line_text_variants(line: dict) -> tuple[str, str]:
    chars = []
    for span in line.get("spans", []):
        size = float(span.get("size") or 0)
        for char in span.get("chars", []):
            value = str(char.get("c") or "")
            if value:
                chars.append({"value": value, "origin": char.get("origin"), "size": size})

    source = "".join(char["value"] for char in chars)
    if " " not in source or len(chars) < 3:
        return source, source

    direction = line.get("dir") or (1.0, 0.0)
    try:
        direction_x = float(direction[0])
        direction_y = float(direction[1])
        direction_length = (direction_x**2 + direction_y**2) ** 0.5
    except (TypeError, ValueError, IndexError):
        return source, source
    if direction_length <= 0:
        return source, source
    direction_x /= direction_length
    direction_y /= direction_length

    remove_indexes = set()
    for index, char in enumerate(chars):
        if char["value"] != " " or not any(item["value"].strip() for item in chars[:index]):
            continue
        next_char = next((item for item in chars[index + 1 :] if item["value"].strip()), None)
        if next_char is None:
            continue
        origin = char.get("origin")
        next_origin = next_char.get("origin")
        if not origin or not next_origin:
            continue
        try:
            advance = (
                (float(next_origin[0]) - float(origin[0])) * direction_x
                + (float(next_origin[1]) - float(origin[1])) * direction_y
            )
        except (TypeError, ValueError, IndexError):
            continue
        tolerance = max(0.25, float(char.get("size") or 0) * 0.03)
        if abs(advance) <= tolerance:
            remove_indexes.add(index)

    normalized = "".join(char["value"] for index, char in enumerate(chars) if index not in remove_indexes)
    return source, normalized


def _read_text(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentReadError("无法识别文本编码，请使用 UTF-8 文档")


def _extract_html(path: Path) -> str:
    html = _read_text(path)
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text("\n", strip=True)


def _extract_openpyxl_workbook(path: Path) -> str:
    value_workbook = load_workbook(path, read_only=True, data_only=True)
    formula_workbook = load_workbook(path, read_only=True, data_only=False)
    try:
        parts = []
        formula_sheets = {sheet.title: sheet for sheet in formula_workbook.worksheets}
        for sheet in value_workbook.worksheets:
            formula_sheet = formula_sheets.get(sheet.title)
            rows = _openpyxl_sheet_rows_text(sheet, formula_sheet)
            if rows:
                parts.append(f"# 工作表：{sheet.title}\n" + "\n".join(rows))
        return "\n\n".join(parts)
    finally:
        value_workbook.close()
        formula_workbook.close()


def _extract_xls(path: Path) -> str:
    workbook = xlrd.open_workbook(str(path), on_demand=True)
    try:
        parts = []
        for sheet in workbook.sheets():
            rows = []
            for row_index in range(sheet.nrows):
                values = [sheet.cell_value(row_index, column_index) for column_index in range(sheet.ncols)]
                row_text = _spreadsheet_row_text(values)
                if row_text:
                    rows.append(row_text)
            if rows:
                parts.append(f"# 工作表：{sheet.name}\n" + "\n".join(rows))
        return "\n\n".join(parts)
    finally:
        workbook.release_resources()


def _openpyxl_sheet_rows_text(sheet, formula_sheet) -> list[str]:
    rows = []
    max_row = max(sheet.max_row or 0, getattr(formula_sheet, "max_row", 0) or 0)
    max_column = max(sheet.max_column or 0, getattr(formula_sheet, "max_column", 0) or 0)
    for row_index in range(1, max_row + 1):
        values = []
        for column_index in range(1, max_column + 1):
            value = sheet.cell(row_index, column_index).value
            if value is None and formula_sheet is not None:
                formula_value = formula_sheet.cell(row_index, column_index).value
                if isinstance(formula_value, str) and formula_value.startswith("="):
                    value = formula_value
            values.append(value)
        row_text = _spreadsheet_row_text(values)
        if row_text:
            rows.append(row_text)
    return rows


def _spreadsheet_rows_text(rows) -> list[str]:
    result = []
    for row in rows:
        row_text = _spreadsheet_row_text(row)
        if row_text:
            result.append(row_text)
    return result


def _spreadsheet_row_text(values) -> str:
    cells = [_spreadsheet_cell_text(value) for value in values]
    while cells and not cells[-1]:
        cells.pop()
    return " | ".join(cells)


def _spreadsheet_cell_text(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()
